import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from utils import generate_random_bytes, generate_random_text, create_generated_folder
from fpdf import FPDF
from docx import Document
from io import BytesIO


def generate_txt(file_path, size_kb):
    """
    Generate a plain text file with random data.
    
    Args:
        file_path (str): Path where the file will be saved
        size_kb (int): Target file size in kilobytes
    """
    # Create generated folder and get full path
    generated_folder = create_generated_folder()
    full_path = os.path.join(generated_folder, file_path)
    
    # Generate random bytes
    random_data = generate_random_bytes(size_kb)
    
    # Write to file
    with open(full_path, 'wb') as f:
        f.write(random_data)
    
    print(f"Successfully generated TXT file: {full_path} ({size_kb} KB)")


def generate_pdf(file_path, size_kb):
    """
    Generate a PDF file with random text content.
    
    Args:
        file_path (str): Path where the file will be saved
        size_kb (int): Target file size in kilobytes
    """
    # Create generated folder and get full path
    generated_folder = create_generated_folder()
    full_path = os.path.join(generated_folder, file_path)
    
    try:
        # Create PDF object
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Helvetica", size=12)
        
        # Add title
        pdf.cell(0, 10, "Random PDF Content", ln=True)
        pdf.ln(10)
        
        # Target size in bytes
        target_bytes = size_kb * 1024
        
        # Add random text content
        # We'll use a simpler approach - just add reasonable chunks of text
        for i in range(min(size_kb // 10, 100)):  # Limit the number of chunks
            # Generate a reasonable chunk of text
            chunk = generate_random_text(200)
            
            # Add to PDF
            pdf.set_font("Helvetica", size=11)
            pdf.multi_cell(0, 7, chunk)
            pdf.ln(3)
            
            # Add new page if needed
            if pdf.get_y() > 270:
                pdf.add_page()
                pdf.set_font("Helvetica", size=12)
                pdf.cell(0, 10, f"Page {pdf.page_no()}", ln=True)
                pdf.ln(5)
        
        # Save the PDF
        pdf.output(full_path)
        
        # Verify the file size
        actual_size = os.path.getsize(full_path)
        print(f"Successfully generated PDF file: {full_path} ({actual_size / 1024:.2f} KB)")
        
    except Exception as e:
        print(f"Error generating PDF file: {e}")
        raise


def generate_docx(file_path, size_kb):
    """
    Generate a DOCX file with random text content.
    
    Args:
        file_path (str): Path where the file will be saved
        size_kb (int): Target file size in kilobytes
    """
    # Create generated folder and get full path
    generated_folder = create_generated_folder()
    full_path = os.path.join(generated_folder, file_path)
    
    try:
        # Create a new Document
        doc = Document()
        
        # Target size in bytes
        target_bytes = size_kb * 1024
        
        # Add title
        title = doc.add_heading('Random Document Content', level=1)
        
        # Add random text until we reach the target size
        current_size = 0
        
        while current_size < target_bytes * 0.7:  # Stop at 70% to account for DOCX overhead
            # Generate a chunk of random text
            chunk_size = min(2000, (target_bytes - current_size) // 8)
            random_chunk = generate_random_text(chunk_size)
            
            # Add paragraph with random text
            paragraph = doc.add_paragraph(random_chunk)
            
            # Check current size by saving to buffer
            buffer = BytesIO()
            doc.save(buffer)
            current_size = len(buffer.getvalue())
        
        # Save the document
        doc.save(full_path)
        
        # Verify the file size
        actual_size = os.path.getsize(full_path)
        print(f"Successfully generated DOCX file: {full_path} ({actual_size / 1024:.2f} KB)")
        
    except Exception as e:
        print(f"Error generating DOCX file: {e}")
        raise